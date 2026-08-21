from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).with_name('Volcre_User_Manual.docx')
NAVY = '0B2545'; BLUE = '2E74B5'; DARK_BLUE = '1F4D78'; MUTED = '667085'; LIGHT = 'E8EEF5'; PALE = 'F4F6F9'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for name, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{name}'))
        if node is None: node = OxmlElement(f'w:{name}'); tcMar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'),'dxa')

def fixed_table(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    ind = OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]/1440); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_font(run, size=11, bold=False, color='000000', italic=False):
    run.font.name = 'Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)

def add_text(doc, text, style=None, before=None, after=None, align=None, size=None, bold=None, color=None, italic=None):
    p = doc.add_paragraph(style=style)
    if before is not None: p.paragraph_format.space_before = Pt(before)
    if after is not None: p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    r=p.add_run(text)
    if size is not None: set_font(r, size, bool(bold), color or '000000', bool(italic))
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25; p.add_run(text); return p

def step(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25; p.add_run(text); return p

def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name, size, color, before, after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK_BLUE,8,4)]:
    st=styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

# Header/footer: editorial_cover pattern adapted for a reference manual
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('VOLCRE | USER MANUAL'); set_font(r,9,True,MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('Volcre Volunteer Management System'); set_font(r,9,False,MUTED)

# Cover
add_text(doc, 'SYSTEM GUIDE', after=18, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, color=BLUE)
add_text(doc, 'Volcre', after=8, align=WD_ALIGN_PARAGRAPH.CENTER, size=30, bold=True, color=NAVY)
add_text(doc, 'User Manual', after=16, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, color=DARK_BLUE)
add_text(doc, 'A role-based guide for administrators, partner organizations, and volunteers', after=110, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=MUTED, italic=True)
add_text(doc, 'Version 1.0 | July 2026', after=4, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, color=NAVY)
add_text(doc, 'Prepared for day-to-day program coordination, project delivery, and reporting.', align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=MUTED)
doc.add_page_break()

heading(doc,'Welcome',1)
add_text(doc,'Volcre is a volunteer-management system for coordinating community programs, partner organizations, projects, volunteers, reports, messaging, and impact data. The experience changes by role: administrators use the web portal; volunteers and partners use the mobile experience.')
heading(doc,'Before you begin',2)
for t in ['Use a supported web browser for the administrator portal.', 'Use the mobile experience for volunteer and partner access. Admin accounts are intentionally limited to the web portal.', 'Keep your email/phone and password current in Profile. New volunteer and partner accounts must be approved before login is unlocked.', 'Use the notification badge and Messages screen to check updates, requests, and conversation activity.'] : bullet(doc,t)
heading(doc,'Roles at a glance',2)
t=doc.add_table(rows=1, cols=3); fixed_table(t,[1800,3000,4560]); hdr=t.rows[0].cells
for cell,text in zip(hdr,['Role','Primary workspace','Typical actions']): set_cell_shading(cell,LIGHT); p=cell.paragraphs[0]; rr=p.add_run(text); set_font(rr,10,True,NAVY)
for row in [('Administrator','Web portal','Approve access; manage projects, partners and volunteers; review reports; monitor analytics.'),('Partner organization','Mobile portal','Maintain organization profile; manage programs; propose or join projects; submit reports.'),('Volunteer','Mobile portal','Maintain skills and availability; join projects; complete tasks; log activity and submit reports.')]:
    cells=t.add_row().cells
    for c,txt in zip(cells,row): c.text=txt

heading(doc,'Contents',2)
for x in ['1. Getting started','2. Administrator guide','3. Partner organization guide','4. Volunteer guide','5. Shared tools: messages, map, reports, profile','6. Statuses, good practice, and support']:
    bullet(doc,x)
doc.add_page_break()

heading(doc,'1. Getting started',1)
heading(doc,'Sign in',2)
for x in ['Open the Volcre sign-in screen.', 'Choose the appropriate portal or mobile role where prompted.', 'Enter your approved email, email username, or phone number and password.', 'Select Sign In. If access is pending or rejected, read the on-screen message before contacting the admin team.'] : step(doc,x)
heading(doc,'Register a new account',2)
add_text(doc,'Registration is available for volunteers and partner organizations. Choose the appropriate role, complete the requested identity and contact fields, and submit the application. Volunteer registration supports profile details such as skills, availability, affiliations, and certifications. Partner registration collects organization and DSWD-related information. Both are reviewed by an administrator before the account can sign in.')
heading(doc,'Navigation',2)
add_text(doc,'Use the sidebar on the administrator web portal or the bottom navigation bar in the mobile portals. The dashboard is the best starting point: it shows current workload, prompts, and shortcuts. Badges on Messages, Reports, Projects, or Users indicate attention is needed.')

heading(doc,'2. Administrator guide',1)
add_text(doc,'Administrators coordinate the full system from the web portal. Use the sidebar to move between Dashboard, Projects, Partners, Volunteers, Map, Messages, Reports, Analytics, Users, and Profile.')
heading(doc,'Dashboard',2)
for x in ['Review summary cards for projects, events, partner status, volunteer activity, and pending work.', 'Use Add Project to begin a new project record.', 'Open dashboard shortcuts to investigate upcoming events, pending reports, messages, or requests.'] : bullet(doc,x)
heading(doc,'User, volunteer, and partner approvals',2)
add_text(doc,'Open Users to review new account requests. Verify the application details, then approve or reject as appropriate. A rejection should include a clear reason. Use Partner Management or the partner approval view to review organization data such as sector, contacts, address, DSWD accreditation, and verification notes. Approval unlocks the relevant portal for the applicant.')
heading(doc,'Projects and lifecycle',2)
for x in ['Open Projects and select Add Project or open an existing project.', 'Enter or review the project title, program module, description, schedule, location, volunteer requirements, and partner information.', 'Use the project lifecycle and planning views to coordinate the project from proposal/planning through execution and completion.', 'Review volunteer requests and assignments; keep dates, capacity, roster, and status current.', 'Use the project calendar or timeline cards to check milestones and event timing.'] : step(doc,x)
heading(doc,'Review proposals and reports',2)
add_text(doc,'Partner proposals can arrive through the Messages workspace. Open the proposal card to review the organization, program module, dates, location, slots, description, community need, deliverables, skills, photo, and attachments. Approve to create a project, or reject with a reason. In Reports, review submitted partner or volunteer reports and use the available report detail/download controls to validate evidence and follow up.')
heading(doc,'Analytics and map',2)
add_text(doc,'Use Analytics to monitor program and activity trends. Use the Map to inspect project and impact locations; select markers to open contextual project information. Treat dashboard and analytics figures as operational indicators and verify underlying project records before external reporting.')
doc.add_page_break()

heading(doc,'3. Partner organization guide',1)
add_text(doc,'Partner organizations use the mobile portal to manage organization information, programs, project participation, proposals, reports, and communication.')
heading(doc,'Partner Dashboard',2)
add_text(doc,'Start on the Partner Dashboard to see active work, program context, messages, and shortcuts. Select a card to open its linked program, project, report, or conversation.')
heading(doc,'Program Management',2)
for x in ['Open Programs to view the organization’s programs and linked projects.', 'Create or update program details only when you have authority to do so.', 'Use the project or proposal entry point when a community need requires a new project for administrator review.', 'Keep objectives, dates, target location, and required resources specific enough for reviewers and volunteers.'] : step(doc,x)
heading(doc,'Propose a project',2)
add_text(doc,'From the proposal flow in Messages or Program Management, complete the project title, detailed description, start/end dates, target location, volunteer slots, community need, expected deliverables, and needed skills. Add a proposal photo or attachments where useful, then submit for review. Track the proposal card in Messages. Approved proposals create a project; rejected proposals show the reviewer’s reason.')
heading(doc,'My Projects',2)
for x in ['Open Projects to view the organization’s current project relationships.', 'Open a project to inspect schedule, location, status, and participation details.', 'Coordinate questions, updates, and supporting evidence through the linked Messages and Reports workflows.', 'Update project information promptly when dates, locations, capacity, or delivery plans change.'] : bullet(doc,x)
heading(doc,'Partner reports',2)
add_text(doc,'Open Reports to submit and review organization reporting for a project. Choose the correct project, complete the report information, attach available evidence where prompted, and submit. Check report status and respond to administrator follow-up in Messages.')

heading(doc,'4. Volunteer guide',1)
add_text(doc,'Volunteers use the mobile portal to discover opportunities, join projects, complete assigned work, track impact, communicate with coordinators, and maintain their profile.')
heading(doc,'Set up your profile',2)
for x in ['Open Profile and select Edit Profile.', 'Keep your name, email/phone, photo, skills, and availability status accurate.', 'Add certifications, trainings, and affiliations when relevant to the opportunities you want to support.', 'Use Change Password when needed; enter matching passwords before saving.'] : step(doc,x)
heading(doc,'Find and join a project',2)
for x in ['Open Projects to browse available programs, projects, and events.', 'Open a project card to review its purpose, dates, location, requirements, capacity, and skills.', 'Select the join/request action. Your request may show as Requested until it is confirmed.', 'Return to the project or dashboard to check your participation and event information.'] : step(doc,x)
heading(doc,'Tasks and activity',2)
add_text(doc,'Open My Tasks to check assigned work and its current status. Complete tasks by their due dates and keep coordinators informed if you cannot finish work. Your dashboard and Profile summarize joined and completed events. Where time logging is offered, submit accurate time-in/time-out or hour details against the correct project.')
heading(doc,'Volunteer reports',2)
add_text(doc,'Open My Reports to view reporting history or create a report for a project. Select the correct project, add a clear activity summary and evidence when requested, then submit. Use Messages to answer any follow-up.')
doc.add_page_break()

heading(doc,'5. Shared tools',1)
heading(doc,'Messages and group chats',2)
for x in ['Open Messages and choose a conversation, contact, project thread, or Event GC (group chat).', 'Type a concise message and send it. Attach a photo or file only when it directly supports coordination or reporting.', 'Use the conversation menu to view members, open event details, or manage group-chat membership when those controls are available to your role.', 'Messages may contain proposal cards. Partners can monitor their submission; administrators can open, approve, or reject it.', 'Unread counts clear after you open the relevant messages or notification area.'] : bullet(doc,x)
heading(doc,'Impact Map',2)
add_text(doc,'The Impact Map is available to all operational roles. Pan and zoom to explore project locations. Select an available marker to inspect its project context. Location visibility depends on project information entered by coordinators.')
heading(doc,'Reports and attachments',2)
add_text(doc,'Use the Reports tab assigned to your role. Attachments may include images or documents. Confirm that the file opens correctly before submission, use meaningful filenames, and never upload sensitive information that is not required for the project.')
heading(doc,'Profile and logout',2)
add_text(doc,'Profile displays role-specific account information. Select Edit Profile to update approved fields and Save to apply changes. Select Logout when using a shared device or when you are finished. You will be asked to confirm before being signed out.')

heading(doc,'6. Statuses and good practice',1)
heading(doc,'Common statuses',2)
t=doc.add_table(rows=1, cols=2); fixed_table(t,[2500,6860])
for cell,text in zip(t.rows[0].cells,['Status','What it means / what to do']): set_cell_shading(cell,LIGHT); r=cell.paragraphs[0].add_run(text); set_font(r,10,True,NAVY)
for row in [('Pending','An account, request, proposal, or report is awaiting a decision or review.'),('Requested','A volunteer has asked to join a project; monitor for administrator/coordinator action.'),('Approved','The account, partner, proposal, or requested action has been accepted. Continue with the linked workflow.'),('Rejected','The action was declined. Read the supplied reason, correct issues where possible, and contact the coordinator if clarification is needed.'),('Submitted','A report or proposal was sent for review. Avoid duplicate submissions unless instructed.')]:
    cells=t.add_row().cells
    for c,txt in zip(cells,row): c.text=txt
heading(doc,'Good operating practice',2)
for x in ['Keep project dates, locations, capacity, contact details, and status current.', 'Use clear titles, short descriptions, and specific deliverables so others can act without guessing.', 'Check Messages and report requests regularly; reply in the related project conversation whenever possible.', 'Protect credentials and log out on shared devices.', 'Do not create duplicate projects, proposals, or reports. Search/open existing records first.'] : bullet(doc,x)
heading(doc,'Troubleshooting',2)
for x in ['Cannot sign in: confirm that you are using the correct role experience, account identifier, and password. New volunteer and partner accounts require approval.', 'No data or request failed: check your connection, refresh/reopen the screen, and try again. If the issue continues, provide the error message and affected project to an administrator.', 'Cannot find a project or message: use the correct role portal and verify that you have been added to the project or conversation.', 'Rejected application: read the reason shown at login or in the relevant record, then contact the admin team with the requested correction.'] : bullet(doc,x)
heading(doc,'Quick handoff checklist',2)
add_text(doc,'Before closing a project cycle: confirm dates and status; verify volunteer participation and task completion; submit required reports with evidence; review outstanding messages; and ensure the project record reflects final outcomes. This keeps dashboards, maps, analytics, and historical records dependable.')

doc.save(OUT)
print(OUT)
